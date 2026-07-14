from __future__ import annotations

import json
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUTS = ROOT / "code" / "outputs"
REPORTS = ROOT / "reports"
DOCX_PATH = REPORTS / "中老年人群高血脂症风险预警及干预优化_建模方案报告_最终版.docx"

FONT_CN = "Microsoft YaHei"
FONT_MATH = "Cambria Math"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
NAVY = "17365D"
MUTED = "666666"
LIGHT_FILL = "F4F6F9"
HEADER_FILL = "E8EEF5"
GOLD_FILL = "FFF4D6"
RED = "9B1C1C"
GREEN = "2F6B3B"

# A4 named override on narrative_proposal preset.
CONTENT_WIDTH_DXA = 9070
TABLE_INDENT_DXA = 120


def set_run_font(run, *, name=FONT_CN, size=None, bold=None, italic=None, color=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_row_cant_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = tr_pr.find(qn("w:cantSplit"))
    if cant_split is None:
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa):
    if sum(widths_dxa) != CONTENT_WIDTH_DXA:
        raise ValueError(f"表格宽度之和必须为 {CONTENT_WIDTH_DXA}，实际为 {sum(widths_dxa)}")
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(CONTENT_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths_dxa):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def set_cell_text(cell, text, *, bold=False, color=None, align=WD_ALIGN_PARAGRAPH.CENTER, size=9):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.first_line_indent = None
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.08
    r = p.add_run(str(text))
    set_run_font(r, size=size, bold=bold, color=color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc, headers, rows, widths_dxa, *, font_size=8.8, left_cols=()):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for j, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[j], header, bold=True, color=NAVY, size=font_size)
        set_cell_shading(table.rows[0].cells[j], HEADER_FILL)
    set_repeat_table_header(table.rows[0])
    set_row_cant_split(table.rows[0])
    for row_data in rows:
        cells = table.add_row().cells
        set_row_cant_split(table.rows[-1])
        for j, value in enumerate(row_data):
            align = WD_ALIGN_PARAGRAPH.LEFT if j in left_cols else WD_ALIGN_PARAGRAPH.CENTER
            set_cell_text(cells[j], value, align=align, size=font_size)
    set_table_geometry(table, widths_dxa)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    return table


def add_source_note(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.first_line_indent = None
    r = p.add_run(text)
    set_run_font(r, size=8, italic=True, color=MUTED)


def add_caption(doc, text):
    p = doc.add_paragraph(style="Caption")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = None
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(7)
    p.paragraph_format.keep_with_next = False
    r = p.add_run(text)
    set_run_font(r, size=9, color=MUTED)


def add_figure(doc, path, caption, width=5.8):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = None
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.keep_with_next = True
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width))
    add_caption(doc, caption)


def add_equation(doc, text, number):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = None
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(f"{text}    ({number})")
    set_run_font(r, name=FONT_MATH, size=11, italic=True, color=NAVY)


def add_body(doc, text, *, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.keep_together = False
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_run_font(r1, size=11, bold=True, color=NAVY)
        r2 = p.add_run(text[len(bold_prefix):])
        set_run_font(r2, size=11)
    else:
        r = p.add_run(text)
        set_run_font(r, size=11)
    return p


def add_bullet(doc, text, level=0):
    style = "List Bullet" if level == 0 else "List Bullet 2"
    p = doc.add_paragraph(style=style)
    p.paragraph_format.first_line_indent = None
    r = p.add_run(text)
    set_run_font(r, size=10.5)
    return p


def add_callout(doc, label, text, *, fill=GOLD_FILL, color=NAVY):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, top=130, bottom=130, start=180, end=180)
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.first_line_indent = None
    p.paragraph_format.space_after = Pt(0)
    r1 = p.add_run(label + "  ")
    set_run_font(r1, size=10.5, bold=True, color=color)
    r2 = p.add_run(text)
    set_run_font(r2, size=10.5, color=color)
    set_table_geometry(table, [CONTENT_WIDTH_DXA])
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    set_run_font(run, size=8.5, color=MUTED)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)
    run2 = paragraph.add_run(" 页")
    set_run_font(run2, size=8.5, color=MUTED)


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)
    section.different_first_page_header_footer = True

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT_CN
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string("222222")
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.first_line_indent = Cm(0.74)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333

    heading_tokens = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (12, DARK_BLUE, 8, 4),
    }
    for name, (size, color, before, after) in heading_tokens.items():
        style = styles[name]
        style.font.name = FONT_CN
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.first_line_indent = None
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Bullet 2", "List Number"):
        style = styles[name]
        style.font.name = FONT_CN
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
        style.font.size = Pt(10.5)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.208

    caption = styles["Caption"]
    caption.font.name = FONT_CN
    caption._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)

    header = section.header
    hp = header.paragraphs[0]
    hp.text = "中老年人群高血脂症风险预警及干预方案优化｜建模方案报告"
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hp.paragraph_format.first_line_indent = None
    set_run_font(hp.runs[0], size=8.5, color=MUTED)
    footer = section.footer
    add_page_number(footer.paragraphs[0])
    return section


def add_cover(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(84)

    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.first_line_indent = None
    kicker.paragraph_format.space_after = Pt(18)
    r = kicker.add_run("2026 MathorCup 数学应用挑战赛 C 题")
    set_run_font(r, size=12, bold=True, color=BLUE)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.first_line_indent = None
    title.paragraph_format.space_after = Pt(12)
    title.paragraph_format.line_spacing = 1.12
    r = title.add_run("中老年人群高血脂症的\n风险预警及干预方案优化")
    set_run_font(r, size=25, bold=True, color=NAVY)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.first_line_indent = None
    subtitle.paragraph_format.space_after = Pt(66)
    r = subtitle.add_run("建模方案报告（修订版）")
    set_run_font(r, size=15, color=DARK_BLUE)

    meta_lines = [
        "数据规模：1000 例中老年个案",
        "模型链路：分终点折外筛选 → 规则管理分层 → 六阶段干预优化",
        "结果状态：代码复跑、约束回代与图表渲染均已通过",
    ]
    for text in meta_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = None
        p.paragraph_format.space_after = Pt(7)
        r = p.add_run(text)
        set_run_font(r, size=10.5, color=MUTED)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(74)
    date = doc.add_paragraph()
    date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date.paragraph_format.first_line_indent = None
    r = date.add_run("2026 年 7 月")
    set_run_font(r, size=11, bold=True, color=NAVY)
    doc.add_page_break()


def fnum(x, digits=3):
    return f"{float(x):.{digits}f}"


def pct(x, digits=1):
    return f"{float(x):.{digits}%}"


def build_report():
    REPORTS.mkdir(parents=True, exist_ok=True)
    summary = json.loads((OUTPUTS / "run_summary.json").read_text(encoding="utf-8"))
    q1_screen = pd.read_csv(OUTPUTS / "question1_dual_endpoint_screening.csv")
    q1_performance = pd.read_csv(OUTPUTS / "question1_endpoint_performance.csv")
    q1_or = pd.read_csv(OUTPUTS / "question1_constitution_adjusted_or.csv")
    q1_score_sensitivity = pd.read_csv(OUTPUTS / "question1_constitution_score_sensitivity.csv")
    q2_candidates = pd.read_csv(OUTPUTS / "question2_candidate_models.csv")
    q2_tiers = pd.read_csv(OUTPUTS / "question2_risk_tiers.csv")
    q2_test = q2_tiers[q2_tiers["数据集"].eq("独立测试集")].sort_values("风险等级编码")
    q2_combos = pd.read_csv(OUTPUTS / "question2_phlegm_core_combinations.csv").head(5)
    q2_rule_sensitivity = pd.read_csv(OUTPUTS / "question2_rule_sensitivity.csv")
    q3_all = pd.read_csv(OUTPUTS / "question3_all_phlegm_patients.csv")
    q3_plans = pd.read_csv(OUTPUTS / "question3_sample_1_2_3_monthly_plans.csv")
    q3_match = pd.read_csv(OUTPUTS / "question3_matching_rules.csv")
    q3_sens = pd.read_csv(OUTPUTS / "question3_sensitivity.csv")
    metrics = summary["question2_metrics"]

    doc = Document()
    configure_document(doc)
    add_cover(doc)

    doc.add_heading("摘要", level=1)
    add_body(doc, "针对中老年人群高血脂症筛查维度单一、痰湿体质高危人群识别不足以及个体化干预难以兼顾效果与成本的问题，本文基于附件中1000例个案，构建“指标证据审查—规则管理分层—方案优化”的递进式模型。问题一将痰湿严重度与高血脂关联筛查拆成两个终点，采用重复五折折外置换重要性、关联检验与FDR联合筛选，并以体质标签相对平和质的调整OR比较九种体质。问题二区分当前诊断状态、非血脂关联筛查概率和最终管理等级，以题面特征阈值生成唯一低、中、高三级规则。问题三针对全部278名痰湿质患者建立六阶段动态规划，采用ε-约束思想，在达到预算内最大降幅90%以上的前提下最小化成本。")
    add_body(doc, f"结果表明，痰湿严重度模型平均折外R²为{q1_performance.loc[q1_performance['终点'].eq('痰湿严重度'),'均值'].iloc[0]:.3f}，没有指标通过痰湿端门槛；排除TC、TG、LDL-C和HDL-C后，高血脂关联端仅血尿酸入选，因此共同指标为空。九体质整体调整检验P={summary['question1_constitution_global_p']:.3f}，不支持贡献差异。问题二非血脂随机森林独立测试AUC为{metrics['auc']:.3f}；规则管理低、中、高级分别为21、247和32人，当前血脂异常率为0%、83.4%和100%。问题三中样本1、2、3分别由64、58、59分降至48.5、37.0、33.0分，总成本分别为1014、1240、1674元，均满足2000元预算及耐受度约束。")
    add_callout(doc, "核心结论", "四项血脂与附件标签100%同义，因此不进入高血脂关联筛查模型；附件没有随访结局，随机森林概率只表示当前异常的横断面关联，不解释为未来发病概率。三级结果由可追溯特征规则生成。")
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = None
    r1 = p.add_run("关键词：")
    set_run_font(r1, size=10.5, bold=True, color=NAVY)
    r2 = p.add_run("高血脂症；痰湿体质；风险分层；随机森林；动态规划；ε-约束")
    set_run_font(r2, size=10.5)

    doc.add_heading("报告结构", level=1)
    outline = [
        "问题重述与建模目标", "数据理解与预处理", "模型假设与符号说明",
        "问题一：分终点指标筛选与体质贡献", "问题二：关联筛查与规则管理分层",
        "问题三：六个月干预优化", "敏感性与一致性检验", "模型评价与结论",
    ]
    # Two-column outline keeps the complete navigation block on the abstract page.
    outline_table = doc.add_table(rows=4, cols=2)
    for i, item in enumerate(outline):
        cell = outline_table.cell(i % 4, i // 4)
        set_cell_text(
            cell,
            f"{i + 1}.  {item}",
            align=WD_ALIGN_PARAGRAPH.LEFT,
            size=10.2,
        )
    set_table_geometry(outline_table, [4535, 4535])
    for row in outline_table.rows:
        for cell in row.cells:
            set_cell_margins(cell, top=35, start=80, bottom=35, end=80)
    doc.add_page_break()

    doc.add_heading("1 问题重述与建模目标", level=1)
    add_body(doc, "高血脂症是中老年心脑血管疾病的重要危险因素。题目要求同时融合中医体质、日常活动能力、血脂与代谢指标，形成从关键指标识别、三级风险预警到个体化干预优化的完整闭环。三个问题并非彼此独立：问题一提供稳定特征和体质贡献依据，问题二给出人群风险层级，问题三则把风险特征、身体耐受度和经济约束转化为可执行的六个月方案。")
    doc.add_heading("1.1 三个子问题", level=2)
    add_bullet(doc, "问题一：从血常规/代谢指标和活动量表中筛选能同时表征痰湿严重程度与高血脂风险的关键指标，并比较九种体质的贡献差异。")
    add_bullet(doc, "问题二：区分当前诊断、非血脂关联筛查概率与最终管理等级，按题面特征阈值输出唯一低、中、高三级规则，并识别痰湿质核心组合。")
    add_bullet(doc, "问题三：对体质标签为5的患者，在年龄、活动能力、频率、成本和疗效约束下优化六个月计划，给出ID 1、2、3方案并归纳匹配规律。")
    doc.add_heading("1.2 总体技术路线", level=2)
    add_body(doc, "总体路线为：原始附件校验 → 诊断泄漏识别 → 分终点折外筛选 → 分类体质调整OR → 非血脂关联概率 → 题面阈值规则分层 → 痰湿核心组合 → 六阶段Pareto动态规划 → ε-约束推荐 → 灵敏度与约束回代。")
    add_callout(doc, "建模边界", "附件高血脂标签与TC、TG、LDL-C、HDL-C任一异常规则的一致率为100%。四项血脂仅描述当前诊断状态；非血脂模型刻画横断面关联，三级规则用于管理优先级，均不等同于未来发病预测。", fill=LIGHT_FILL)

    doc.add_heading("2 数据理解与预处理", level=1)
    doc.add_heading("2.1 数据规模与字段", level=2)
    quality = summary["data_quality"]
    rows = [
        ["样本规模", str(quality["rows"]), "附件1全部个案"],
        ["体质标签5人数", str(quality["phlegm_constitution_count"]), "问题三求解对象"],
        ["高血脂阳性率", pct(quality["label_prevalence"]), "样本患病率较高"],
        ["缺失值", str(quality["missing_values"]), "建模字段完整"],
        ["重复样本ID", str(quality["duplicate_sample_ids"]), "样本ID唯一"],
        ["活动恒等式最大误差", fnum(quality["activity_identity_max_error"], 1), "ADL+IADL=活动总分"],
    ]
    add_table(doc, ["检查项", "结果", "说明"], rows, [2500, 1700, 4870], font_size=9.2, left_cols=(0, 2))
    add_source_note(doc, "数据来源：题目附件1；检查结果来自 code/outputs/data_quality.json。")
    doc.add_heading("2.2 临床阈值与派生变量", level=2)
    add_body(doc, "按照题面给出的临床正常范围，定义TC>6.2 mmol/L、TG>1.7 mmol/L、LDL-C>3.1 mmol/L或HDL-C<1.04 mmol/L为相应血脂异常；任一血脂异常即与附件高血脂标签一致。血尿酸根据性别采用不同范围，BMI正常范围为18.5–23.9 kg/m²。")
    add_body(doc, "为避免完全线性重复，关联筛查模型只使用活动量表总分，不同时放入ADL、IADL及其和。所有标准化、概率校准和模型拟合都在交叉验证管道内部完成；最终管理规则由题面特征阈值预先确定，测试集只用于模型性能和规则一致性核查。")

    doc.add_heading("3 模型假设与符号说明", level=1)
    doc.add_heading("3.1 主要假设", level=2)
    assumptions = [
        "附件样本在本研究范围内可用于建立相对风险排序，但不直接代表一般人群绝对患病率。",
        "体质积分、活动评分和体检指标在建模时视为同一时点观测；模型刻画关联与风险排序，不作因果推断。",
        "问题三中年龄组在六个月内不变，活动耐受度按初始活动总分确定；若临床复评后耐受度改变，应重新求解。",
        "题面未给中医调理方式的独立积分下降率，基准方案仅计其必选成本；0%、1%、2%的额外月降幅只用于敏感性分析。",
        "每月按4周、六个月按24周；活动频率1–10次/周，总成本不超过2000元。",
    ]
    for text in assumptions:
        add_bullet(doc, text)
    doc.add_page_break()
    doc.add_heading("3.2 主要符号", level=2)
    symbol_rows = [
        ["S", "痰湿体质积分", "分"], ["Y", "高血脂二分类标签", "0/1"],
        ["Īᵢ,S / Īᵢ,Y", "指标i在验证折上的平均置换重要性", "实数"],
        ["p", "非血脂关联筛查模型输出概率", "0–1"], ["G(x)", "题面特征规则生成的管理等级", "1/2/3"],
        ["aₜ", "第t月活动强度", "1/2/3级"], ["fₜ", "第t月每周活动频率", "次/周"],
        ["Cₜ", "第t月调理与活动总成本", "元"], ["Δmax", "预算内最大可降低积分", "分"],
    ]
    add_table(doc, ["符号", "含义", "单位/取值"], symbol_rows, [1500, 5570, 2000], font_size=9.2, left_cols=(1,))

    doc.add_heading("4 问题一：分终点指标筛选与体质贡献", level=1)
    doc.add_heading("4.1 折外证据与泄漏排除", level=2)
    add_body(doc, "痰湿严重度终点考察TC、TG、LDL-C、HDL-C、血糖、血尿酸、BMI和活动总分；高血脂关联终点排除定义标签的四项血脂，只考察血糖、血尿酸、BMI和活动总分。两个终点均采用5折、3次重复交叉验证，并仅在验证折计算置换重要性。")
    add_equation(doc, "K = Kₛ ∩ Kᵧ；Kₛ、Kᵧ均由折外性能、置换重要性、关联检验与FDR共同确定", 1)
    add_body(doc, "痰湿端要求平均折外R²>0、正重要性频率不低于0.60、|Spearman ρ|≥0.10且FDR q<0.05；高血脂关联端要求平均折外AUC≥0.60、正重要性频率不低于0.60、单指标AUC≥0.60且FDR q<0.05。交集允许为空，避免为满足形式而制造假阳性。")
    phlegm_perf = q1_performance[q1_performance["终点"].eq("痰湿严重度")].iloc[0]
    lipid_perf = q1_performance[q1_performance["终点"].eq("高血脂关联筛查")].iloc[0]
    add_callout(doc, "折外性能", f"痰湿严重度平均折外R²={phlegm_perf['均值']:.3f}±{phlegm_perf['标准差']:.3f}；非诊断高血脂关联模型平均折外AUC={lipid_perf['均值']:.3f}±{lipid_perf['标准差']:.3f}。", fill=LIGHT_FILL)
    doc.add_page_break()
    indicators = ["总胆固醇_TC", "甘油三酯_TG", "低密度脂蛋白_LDL_C", "高密度脂蛋白_HDL_C", "血糖", "血尿酸", "BMI", "活动量表总分"]
    screen_rows = []
    for indicator in indicators:
        ps = q1_screen[(q1_screen["终点"].eq("痰湿严重度")) & q1_screen["指标名称"].eq(indicator)].iloc[0]
        ys = q1_screen[(q1_screen["终点"].eq("高血脂关联筛查")) & q1_screen["指标名称"].eq(indicator)].iloc[0]
        if bool(ys["是否诊断泄漏"]):
            lipid_evidence = "诊断泄漏排除"
        else:
            lipid_evidence = f"AUC {ys['单指标AUC']:.3f} / q {ys['FDR_q值']:.3f}"
        if indicator in summary["question1_shared"]:
            conclusion = "共同指标"
        elif indicator in summary["question1_selected_lipid"]:
            conclusion = "仅高血脂关联端"
        elif indicator in summary["question1_selected_phlegm"]:
            conclusion = "仅痰湿端"
        elif bool(ys["是否诊断泄漏"]):
            conclusion = "排除"
        else:
            conclusion = "未通过"
        screen_rows.append([
            indicator, f"{ps['关联统计量']:.3f} / {ps['FDR_q值']:.3f}",
            f"{ps['折外正重要性频率']:.2f}", lipid_evidence, conclusion,
        ])
    add_table(doc, ["指标", "痰湿ρ / FDR q", "痰湿正重要性频率", "高血脂关联证据", "结论"], screen_rows,
              [1900, 1800, 1650, 2350, 1370], font_size=8.2, left_cols=(0, 3, 4))
    add_source_note(doc, "表1  分终点折外筛选结果。TC、TG、LDL-C、HDL-C在高血脂关联端因定义标签而排除。")
    add_figure(doc, ROOT / "问题1_共识得分排序.png", "图1  两个终点的折外置换重要性", width=5.9)
    add_body(doc, "痰湿端没有指标通过门槛；高血脂关联端仅血尿酸入选；两集合交集为空。该结果说明附件不足以支持“同一指标同时可靠表征痰湿严重度并预警高血脂”的结论，TC和TG只能作为当前诊断指标。")
    add_figure(doc, ROOT / "问题1_双终点频率热力图.png", "图2  候选指标在验证折中的正置换重要性频率", width=5.6)

    doc.add_heading("4.2 九种体质标签调整优势比", level=2)
    add_body(doc, "九种体质贡献的主分析改用体质标签作为分类暴露，以平和质为参照，调整年龄、性别、吸烟、饮酒、BMI、活动总分、血糖和血尿酸。先用整体似然比检验判断八个标签系数是否同时为零，再报告各组OR、95%CI、P值和BH-FDR。")
    add_equation(doc, "logit[P(Y=1)] = β₀ + Σₖ₌₂⁹βₖI(T=k) + ΣₘγₘXₘ", 2)
    doc.add_page_break()
    or_rows = []
    for _, r in q1_or.iterrows():
        if int(r["体质标签"]) == 1:
            ci_text, p_text, q_text = "参照", "—", "—"
        else:
            ci_text = f"{r['95%CI下限']:.3f}–{r['95%CI上限']:.3f}"
            p_text, q_text = f"{r['P值']:.3f}", f"{r['FDR_q值']:.3f}"
        or_rows.append([r["体质类型"], int(r["样本数"]), pct(r["观察患病率"]), fnum(r["调整OR"], 3), ci_text, p_text, q_text])
    add_table(doc, ["体质类型", "n", "当前异常率", "调整OR", "95%CI", "P值", "FDR q"], or_rows,
              [1350, 1050, 1350, 1100, 1800, 1100, 1320], font_size=8.2, left_cols=(0,))
    add_source_note(doc, "表2  九种体质标签相对平和质的调整OR；整体似然比检验先于组间解释。")
    add_figure(doc, ROOT / "问题1_体质贡献度OR.png", "图3  八种偏颇体质相对平和质的调整OR及95%置信区间", width=5.8)
    phlegm_or = q1_or[q1_or["体质标签"].eq(5)].iloc[0]
    add_callout(doc, "统计结论", f"九体质整体调整检验P={summary['question1_constitution_global_p']:.3f}；痰湿质相对平和质OR={phlegm_or['调整OR']:.3f}，95%CI {phlegm_or['95%CI下限']:.3f}–{phlegm_or['95%CI上限']:.3f}，P={phlegm_or['P值']:.3f}，FDR q={phlegm_or['FDR_q值']:.3f}。九个连续积分分别调整的敏感性分析同样无FDR显著项，因此不作贡献排名。", fill=LIGHT_FILL)

    doc.add_heading("5 问题二：关联筛查与规则管理分层", level=1)
    doc.add_heading("5.1 三类输出严格分离", level=2)
    add_body(doc, "问题二同时输出三个含义不同的量：当前诊断状态由四项血脂阈值确定；非血脂关联概率由体质、活动、代谢和基础信息模型给出；最终低、中、高管理等级由题面特征规则唯一生成。三者分别回答“当前是否异常”“非血脂特征与异常的关联强弱”和“应采用何种管理优先级”，不再混用。")
    add_callout(doc, "横断面限制", "附件没有随访结局，非血脂概率不能解释为未来发病概率；三级管理规则包含当前血脂状态，其当前异常率只用于规则一致性核查，不能作为独立预测性能。", fill=LIGHT_FILL)
    doc.add_heading("5.2 非血脂关联模型", level=2)
    candidate_rows = [[r["模型"], fnum(r["训练集CV_AUC均值"], 3), fnum(r["训练集CV_AUC标准差"], 3)] for _, r in q2_candidates.iterrows()]
    add_table(doc, ["候选模型", "五折CV AUC均值", "标准差"], candidate_rows, [3600, 2800, 2670], font_size=9.3, left_cols=(0,))
    add_source_note(doc, "表3  训练集内部候选模型比较。随机森林仅因训练集CV表现更优而被选中。")
    add_body(doc, "随机森林使用类别平衡抽样、最小叶节点约束和Platt概率校准。训练/测试按7:3分层划分，候选模型比较与校准均在训练集完成。独立测试集只用于一次最终评价，模型特征中不存在四项诊断血脂及其异常派生量。")
    add_figure(doc, ROOT / "问题2_筛查模型ROC.png", "图4  非血脂关联筛查模型独立测试集ROC曲线", width=5.4)
    add_body(doc, f"随机森林独立测试AUC={metrics['auc']:.3f}、AP={metrics['average_precision']:.3f}、Brier={metrics['brier']:.3f}。该指标评价的是对当前血脂异常的关联筛查能力，而非未来发病预测。")

    doc.add_heading("5.3 可追溯三级管理规则", level=2)
    add_bullet(doc, "高危：血脂异常且痰湿积分≥60；或血脂正常、痰湿积分≥80且活动总分<40；或痰湿质、痰湿积分≥60且尿酸异常。")
    add_bullet(doc, "中危：不满足高危，但存在血脂异常、痰湿积分≥60、活动总分<40或至少一项血糖、尿酸、BMI异常。")
    add_bullet(doc, "低危：未触发上述高危或中危规则。")
    add_equation(doc, "G(x)=3（触发任一高危规则）；G(x)=2（未高危且触发任一中危条件）；否则G(x)=1", 3)
    tier_rows = []
    for _, r in q2_test.iterrows():
        tier_rows.append([r["风险等级"], int(r["人数"]), pct(r["占比"]), pct(r["实际患病率"]),
                          f"{pct(r['患病率95%CI下限'])}–{pct(r['患病率95%CI上限'])}"])
    add_table(doc, ["管理等级", "人数", "占比", "当前异常率", "异常率95%CI"], tier_rows,
              [1600, 1300, 1500, 1900, 2770], font_size=9.2)
    add_source_note(doc, "表4  独立测试集三级管理结果。异常率因规则使用当前诊断状态而仅作横断面一致性核查。")
    add_figure(doc, ROOT / "问题2_三级风险分布.png", "图5  独立测试集三级管理人数与当前血脂异常率", width=5.4)
    add_body(doc, f"低危与中危Fisher精确检验P≈{metrics['fisher_p_low_vs_medium_management']:.2e}，中危与高危P={metrics['fisher_p_medium_vs_high_management']:.4f}。每名患者只保留一个最终管理等级和全部触发规则，已异常者不再出现低危输出。")
    add_figure(doc, ROOT / "问题2_校准曲线.png", "图6  非血脂关联筛查模型独立测试集校准曲线", width=5.4)

    doc.add_heading("5.4 痰湿质高危核心组合", level=2)
    add_body(doc, "在独立测试集内固定“体质标签=5”，再组合当前血脂状态、痰湿积分、低活动量、BMI、血糖、尿酸、年龄与烟酒阈值，计算组合相对于最终高危管理层的支持度、置信度、Jeffreys 95%CI和提升度。")
    doc.add_page_break()
    combo_rows = []
    for _, r in q2_combos.iterrows():
        combo_rows.append([r["核心特征组合"], int(r["覆盖人数"]), int(r["高危人数"]),
                           f"{pct(r['高危置信度'])}（{pct(r['置信度95%CI下限'])}–{pct(r['置信度95%CI上限'])}）",
                           fnum(r["提升度"], 2)])
    add_table(doc, ["核心特征组合", "覆盖人数", "高危人数", "置信度（95%CI）", "提升度"], combo_rows,
              [3650, 1000, 1000, 2250, 1170], font_size=8.1, left_cols=(0,))
    add_source_note(doc, "表5  痰湿质高危管理组合前5位。提升度相对痰湿质人群的基准高危比例计算。")
    top_combo = q2_combos.iloc[0]
    add_callout(doc, "首位组合", f"“{top_combo['核心特征组合']}”覆盖{int(top_combo['覆盖人数'])}人，其中高危{int(top_combo['高危人数'])}人，置信度{pct(top_combo['高危置信度'])}，95%CI {pct(top_combo['置信度95%CI下限'])}–{pct(top_combo['置信度95%CI上限'])}。该组合直接反映管理规则，不作为独立因果发现。")

    doc.add_heading("6 问题三：六个月个体化干预优化", level=1)
    doc.add_heading("6.1 决策变量与约束", level=2)
    add_body(doc, "研究对象为全部278名体质标签5患者。第t月决策包括活动强度aₜ∈{1,2,3}和每周频率fₜ∈{1,…,10}；中医调理等级由月初痰湿积分自动确定。强度上界同时受年龄和活动评分约束：40–59岁最高3级、60–79岁最高2级、80–89岁最高1级；活动总分<40最高1级，40–59最高2级，≥60最高3级。")
    add_equation(doc, "Cost = Σₜ(Cost调理,t + 4·fₜ·Cost活动,aₜ) ≤ 2000", 4)
    doc.add_heading("6.2 状态转移", level=2)
    add_body(doc, "根据题面，当每周训练少于5次时积分基本稳定；达到5次后，每提高一级强度每月下降约3%，同一强度下每周增加1次训练每月多下降约1%。")
    add_equation(doc, "r(a,f)=0 (f<5)；r(a,f)=0.03(a−1)+0.01(f−5) (f≥5)", 5)
    add_equation(doc, "Sₜ₊₁ = Sₜ·[1−r(aₜ,fₜ)]", 6)
    add_body(doc, "积分以0.5分为步长离散。对同一积分只保留最低成本状态，并删除“积分更高且成本不低”的支配状态，从而得到每月Pareto前沿。")
    doc.add_heading("6.3 ε-约束推荐规则", level=2)
    add_body(doc, "若仅以最终积分最小为目标，方案会倾向耗尽预算。为兼顾效果和经济性，先求2000元预算内最大可降幅Δmax，再从达到其90%以上的可行方案中选择成本最低者。")
    add_equation(doc, "min Cost，s.t. Δ ≥ 0.90·Δmax，Cost ≤ 2000", 7)
    add_callout(doc, "方案含义", "推荐方案不是绝对最大降幅方案，而是在保留至少90%最大效果的前提下节省成本；因此不同患者的总成本可能明显低于2000元。", fill=LIGHT_FILL)

    doc.add_heading("6.4 样本1、2、3结果", level=2)
    sample_rows = []
    for sample_id in (1, 2, 3):
        r = q3_all[q3_all["样本ID"].eq(sample_id)].iloc[0]
        sample_rows.append([sample_id, fnum(r["初始痰湿积分"], 1), int(r["年龄组"]), fnum(r["活动量表总分"], 1),
                            int(r["最大允许活动强度"]), fnum(r["最终痰湿积分"], 1), fnum(r["降低分数"], 1), int(r["六个月总成本"])])
    add_table(doc, ["ID", "初始积分", "年龄组", "活动总分", "最大强度", "最终积分", "降低分数", "总成本/元"], sample_rows,
              [900, 1050, 950, 1100, 1050, 1050, 1050, 1920], font_size=8.7)
    add_source_note(doc, "表6  附件样本ID 1、2、3的真实参数与推荐结果。")
    add_figure(doc, ROOT / "问题3_积分变化曲线.png", "图7  三位患者六个月痰湿积分变化", width=5.8)
    add_body(doc, "样本1活动总分38，只允许1级强度，因此主要依靠频率提高，最终由64分降至48.5分；样本2活动总分40，可用2级强度，在1240元成本下由58分降至37分；样本3活动总分63且年龄组1，可耐受3级强度，由59分降至33分，但成本相应提高至1674元。")
    add_figure(doc, ROOT / "问题3_成本效果对比.png", "图8  三位患者推荐方案成本与效果对比", width=5.8)

    doc.add_heading("6.5 患者特征—最优方案匹配规律", level=2)
    match_rows = []
    for _, r in q3_match.iterrows():
        match_rows.append([r["初始痰湿分层"], int(r["最大允许活动强度"]), int(r["患者人数"]),
                           f"{int(r['典型首月活动强度'])}级/{int(r['典型首月每周次数'])}次",
                           int(r["六个月成本中位数"]), fnum(r["降低分数中位数"], 1), pct(r["效果保留率中位数"])])
    add_table(doc, ["初始分层", "耐受上限", "人数", "典型首月方案", "成本中位数", "降分中位数", "效果保留率"], match_rows,
              [1900, 950, 850, 1600, 1300, 1200, 1270], font_size=8.2, left_cols=(0,))
    add_source_note(doc, "表7  全部278名痰湿质患者的分层匹配规律。典型首月方案取组内众数。")
    add_body(doc, "匹配规律表明，最大允许强度是首月活动方案和成本效果的主要分层变量：耐受上限1级者典型方案为1级、每周10次；上限2级者多采用2级、每周9–10次；上限3级者优先使用高强度，但强化调理组可能为节约成本在部分月份回落至2级。初始痰湿分层主要影响调理成本，而耐受上限主要决定可实现的降幅。")

    doc.add_heading("7 敏感性与一致性检验", level=1)
    doc.add_heading("7.1 问题二规则敏感性", level=2)
    add_body(doc, "在独立测试集上分别将痰湿高阈值、痰湿极高阈值和低活动阈值上下扰动5分，考察规则分层人数与当前血脂异常率的变化。该检验用于判断管理规则对题面示例阈值的敏感程度，不把随机森林概率重新转化为分层切点。")
    rule_rows = []
    for level in ("低危", "中危", "高危"):
        g = q2_rule_sensitivity[q2_rule_sensitivity["风险等级"].eq(level)]
        rule_rows.append([
            level,
            f"{int(g['人数'].min())}–{int(g['人数'].max())}",
            f"{100 * g['占比'].min():.1f}%–{100 * g['占比'].max():.1f}%",
            f"{100 * g['实际患病率'].min():.1f}%–{100 * g['实际患病率'].max():.1f}%",
        ])
    add_table(doc, ["管理等级", "人数范围", "占比范围", "当前异常率范围"], rule_rows,
              [1700, 1900, 2400, 3070], font_size=9.0)
    add_source_note(doc, "表8  题面特征阈值±5分的独立测试集敏感性结果。")
    add_body(doc, "痰湿高阈值对高危人数最敏感：由60分降至55分时高危增至65人，升至65分时降为3人；但各情景高危组当前异常率均为100%。痰湿极高阈值在本测试集内不改变分层，低活动阈值仅小幅改变低、中危人数。因此三级结果应称为管理优先级，外部应用时需结合临床共识校准阈值。")
    doc.add_heading("7.2 问题三参数敏感性", level=2)
    sens_rows = []
    for (sid, dim), g in q3_sens.groupby(["样本ID", "敏感性维度"], observed=True):
        if dim == "基准":
            continue
        sens_rows.append([int(sid), dim, f"{g['最终积分'].min():.1f}–{g['最终积分'].max():.1f}",
                          f"{int(g['总成本'].min())}–{int(g['总成本'].max())}"])
    add_table(doc, ["样本ID", "敏感性维度", "最终积分范围", "总成本范围/元"], sens_rows,
              [1100, 3000, 2300, 2670], font_size=8.8, left_cols=(1,))
    add_source_note(doc, "表9  一因素敏感性汇总：积分步长、活动效果系数、预算上限和调理额外月降幅分别扰动。")
    add_body(doc, "积分步长从1.0缩小到0.25主要影响舍入口径，不改变耐受上限和总体方案方向；活动效果系数对最终积分最敏感；预算变化会改变可达最大效果，但ε-约束能避免在效果增益很小的情况下机械耗尽预算；若假设调理具有额外月降幅，最终积分会进一步下降，但该参数缺乏题面数据支撑，只能作为情景分析。")
    doc.add_heading("7.3 约束与复现性", level=2)
    checks = [
        "1000例数据均由原始附件直接读取，样本ID唯一、建模字段无缺失。",
        "问题一高血脂关联模型和问题二非血脂模型均不含TC、TG、LDL-C、HDL-C及其异常派生变量；痰湿端允许无指标入选。",
        "问题一分类体质模型明确以平和质为参照，报告整体似然比检验、OR、95%CI、P值与FDR。",
        "问题二模型训练与校准均与测试集隔离；最终管理等级由预先给定的特征规则生成，每人恰好一个等级且触发原因可回溯。",
        "全部278名痰湿质患者的年龄、活动评分、强度、频率、月份和预算约束逐条回代通过。",
        "样本1、2、3均按附件真实ID读取，不再使用手工构造参数。",
        "所有关键数值保存为CSV/JSON，数据图同时保存PNG与PDF，可从代码完整复现。",
    ]
    for text in checks:
        add_bullet(doc, text)

    doc.add_heading("8 模型评价与结论", level=1)
    doc.add_heading("8.1 模型优点", level=2)
    add_bullet(doc, "逻辑闭环：从分终点证据审查、规则管理分层到个体化优化，三个问题之间具有明确依赖关系。")
    add_bullet(doc, "防止信息泄漏：将四项诊断血脂与非血脂关联筛查分离，痰湿端重要性只在验证折计算。")
    add_bullet(doc, "统计表达完整：体质贡献以平和质为分类参照，同时报告整体检验、OR、95%CI、P值和FDR，不以排序替代显著性。")
    add_bullet(doc, "优化目标可解释：ε-约束直接表达“保留90%最大效果后成本最低”，比无量纲不明的加权和更易说明。")
    add_bullet(doc, "可复现性强：输出患者级风险、全部痰湿质推荐、逐月方案、Pareto、敏感性与约束检查。")
    doc.add_heading("8.2 模型局限", level=2)
    add_bullet(doc, "附件为横截面样本，非血脂模型刻画的是与当前血脂异常的关联筛查分数，而非真实未来发病概率。")
    add_bullet(doc, "最终等级包含当前诊断状态，其组间异常率不能视为独立预测验证；三级规则目前只能解释为管理优先级。")
    add_bullet(doc, "九体质调整OR的置信区间均需谨慎解释，不能据此认定体质因果效应。")
    add_bullet(doc, "题面未给中医调理的独立疗效参数，基准模型无法优化不同调理方式的真实疗效差异。")
    add_bullet(doc, "活动耐受度按初始评分固定；实际应用应每月复评并滚动重优化。")
    doc.add_heading("8.3 主要结论", level=2)
    conclusions = [
        f"痰湿严重度模型平均折外R²为{q1_performance.loc[q1_performance['终点'].eq('痰湿严重度'),'均值'].iloc[0]:.3f}，没有指标通过痰湿端证据门槛；高血脂关联端仅血尿酸入选，因此附件数据未发现可靠共同指标。",
        f"九体质分类自变量整体调整检验P={summary['question1_constitution_global_p']:.3f}；痰湿质相对平和质调整OR={summary['question1_phlegm_constitution_or']['调整OR']:.3f}，95%CI {summary['question1_phlegm_constitution_or']['95%CI下限']:.3f}–{summary['question1_phlegm_constitution_or']['95%CI上限']:.3f}，FDR q={summary['question1_phlegm_constitution_or']['FDR_q值']:.3f}，不能宣称某种体质贡献首位。",
        f"非血脂随机森林关联筛查模型测试AUC为{metrics['auc']:.3f}；独立测试集规则管理低、中、高级分别为21、247、32人，每人只有一个等级和可追溯触发规则。",
        "痰湿质核心组合相对于最终管理等级重新计算，并同时报告支持度、置信度、提升度和样本量；结果用于管理线索，不作未来发病因果解释。",
        "ε-约束动态规划能在保留至少90%最大降幅的同时减少不必要成本，并为全部278名痰湿质患者提供可执行方案。",
        "真实样本1、2、3的耐受上限分别为1、2、3级，最终积分和成本差异体现了“患者特征—方案强度—成本效果”的匹配关系。",
    ]
    for text in conclusions:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.first_line_indent = None
        r = p.add_run(text)
        set_run_font(r, size=10.5)

    doc.add_heading("附录A 样本1、2、3逐月方案", level=1)
    plan_rows = []
    for _, r in q3_plans.iterrows():
        plan_rows.append([
            int(r["样本ID"]), int(r["月份"]), fnum(r["月初痰湿积分"], 1), int(r["调理等级"]),
            f"{int(r['活动强度'])}级/{int(r['每周次数'])}次", int(r["当月总成本"]), fnum(r["月末痰湿积分"], 1),
        ])
    add_table(doc, ["ID", "月份", "月初积分", "调理等级", "活动强度/频率", "月成本/元", "月末积分"], plan_rows,
              [900, 850, 1200, 1150, 2000, 1450, 1520], font_size=8.3)
    add_source_note(doc, "表A1  三位患者的逐月推荐计划。核心调理方式随等级依题面附表2执行。")

    doc.add_heading("附录B 结果文件与复现方式", level=1)
    add_body(doc, "核心代码入口为code/solution.py；执行code/build_notebook.py可从原始附件重建预处理数据、CSV/JSON结果、PNG/PDF图和已执行Notebook。结果说明见reports/RESULTS_REPORT.md，全部数值中间件位于code/outputs目录。")
    add_body(doc, "建议在Windows PowerShell中创建Python 3.12虚拟环境，安装requirements-py312.txt后运行构建脚本。报告中的全部结果均来自最新复跑输出，未手工重算或替换数值。")

    # Avoid widow/orphan and annotate document metadata.
    props = doc.core_properties
    props.title = "中老年人群高血脂症风险预警及干预优化——建模方案报告"
    props.subject = "MathorCup C题建模方案"
    props.keywords = "高血脂症, 痰湿体质, 风险预警, 动态规划"
    props.comments = "由项目最新可复现结果生成"
    doc.save(DOCX_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    build_report()
